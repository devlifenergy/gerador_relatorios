import streamlit as st
import pandas as pd
from docx import Document
import re
import io
import zipfile
import requests
import json
import time
from streamlit_gsheets import GSheetsConnection

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Sistema Integrado DUPAR", page_icon="📊", layout="wide")

# --- CONSTANTES GLOBAIS ---

# 1. ORIGEM (Dados Brutos - Leitura)
URL_DADOS_BRUTOS = "https://docs.google.com/spreadsheets/d/1cvt_tSKEqqH5pvUXVJwAdbCeQ4o3KoBiQMOds7KLeok/edit?usp=sharing"

# 2. DESTINO (Planilha Mestra - Escrita/Leitura para GPT)
SHEET_ID_DESTINO = "1xny0NMN0JrpiC8cxc5b8bL2QrBwAVJJqIG8g1jty2Ps" 
URL_EXPORT_EXCEL = f"https://docs.google.com/spreadsheets/d/{SHEET_ID_DESTINO}/export?format=xlsx"
URL_G_SHEET_LINK = f"https://docs.google.com/spreadsheets/d/{SHEET_ID_DESTINO}/edit?usp=sharing"

# 3. TEMPLATE PADRÃO
TEMPLATE_PADRAO = """Faça um relatório do modelo que estamos treinando com os seguintes itens: identificação (nome, cpf, e-mail, data do nascimento, naturalidade, objetivo de participação), registro de aplicação (identidade assistida ou nome, nome do aplicador, escolha da atividade - instrutor ou a própria pessoa, tempo dele resposta e data de aplicação, fractal de comportamento), registro de dados (fazer um quadro com as seguintes colunas e dados; numero de resposta, respostas, hierarquia e padrão de comportamento psicológico identificado), abaixo do quadro escreva um paragrafo sobre a interpretação do quadro de registro de dados, e a seguir outro paragrafo com as sugestões (habilidades a 
desenvolver fundamentado no paragrafo da interpretação). utilize os dados a seguir: 

relatório 1: {{Horario de inicio}} {{Horario de termino}} {{Tempo de resposta}} {{Data de Aplicação}} {{Nome}} {{E-mail}} {{Naturalidade}} {{CPF}} {{Data de nascimento}} {{Objetivo}} {{Tipo de Aplicação}} {{Nome do Aplicador}}  {{Escolha da Atividade}} fractal de comportamento: 1- "{{Pergunta}}" {{RESPOSTA 1}} {{HIERAQUIA 1}} {{JUSTIFICATIVA 1}} {{RESPOSTA 2}} {{HIERAQUIA 2}} {{JUSTIFICATIVA 2}} {{RESPOSTA 3}} {{HIERAQUIA 3}} {{JUSTIFICATIVA 3}} {{FEEDBACK FINAL}}.

relatório 2: {{Horario de inicio}} {{Horario de termino}} {{Tempo de resposta}} {{Data de Aplicação}} {{Nome}} {{E-mail}} {{Naturalidade}} {{CPF}} {{Data de nascimento}} {{Objetivo}} {{Tipo de Aplicação}} {{Nome do Aplicador}} {{Escolha da Atividade}} fractal de comportamento 2 - "{{Pergunta_2}}" {{RESPOSTA 1_2}} {{HIERAQUIA 1_2}} {{JUSTIFICATIVA 1_2}} {{RESPOSTA 2_2}} {{HIERAQUIA 2_2}} {{JUSTIFICATIVA 2_2}} {{RESPOSTA 3_2}} {{HIERAQUIA 3_2}} {{JUSTIFICATIVA 3_2}} {{FEEDBACK FINAL_2}}.

cruze as informações levantadas e faça uma síntese em um paragrafo, dos padrões de comportamento psicológico recorrentes nas respostas dos três fractais.
em seguida escreva um outro paragrafo de recomendação de desenvolvimento de habilidades.
em seguida combinar as informações e categorizar como padrões de comportamento do usuário tem conteúdos definidos como: ● socialização: atributo relacionado com as interações do usuário com outros indivíduos, sejam familiares, amigos ou colegas de trabalho;
● reflexão: atributo relacionado com a reflexão interior do usuário sobre as suas questões de vida e aspectos maiores do contexto no qual ele habita;
● lazer: atributo relacionado com a realização de atividades que promovem o prazer e a felicidade do usuário, sejam elas ao ar livre ou em casa;
● propósito: atributo relacionado com a motivação pessoal e os objetivos do usuário, ditando suas ambições, perspectivas de futuro e conquistas;
● sentimento: atributo relacionado com o equilíbrio emocional do usuário e sua relação positiva com os aspectos sentimentais internos e externos;
lazer, socialização, reflexão, propósito e sentimento numa métrica de 0 a 100%"""

# ==============================================================================
# FUNÇÕES AUXILIARES (GERADOR GPT)
# ==============================================================================
def ler_e_substituir_template(template_text, dados):
    def substituir(match):
        variavel = match.group(1).strip()
        valor = dados.get(variavel)
        return str(valor).strip() if valor is not None else ""
    return re.sub(r'\{\{([^{}]+)\}\}', substituir, template_text)

def formatar_paragrafo_com_negrito(paragrafo, texto):
    partes = re.split(r'(\*\*.*?\*\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragrafo.add_run(parte[2:-2])
            run.bold = True
        else:
            paragrafo.add_run(parte)

def criar_docx_bytes(texto_resposta):
    doc = Document()
    linhas = texto_resposta.split('\n')
    for linha in linhas:
        linha = linha.strip()
        if not linha: continue
        if linha.startswith('#'):
            nivel = linha.count('#')
            texto_limpo = linha.lstrip('#').strip()
            doc.add_heading(texto_limpo, level=min(nivel, 9))
        elif linha.startswith('- ') or linha.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            texto_limpo = linha[2:]
            formatar_paragrafo_com_negrito(p, texto_limpo)
        else:
            p = doc.add_paragraph()
            formatar_paragrafo_com_negrito(p, linha)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def chamar_gpt(api_key, prompt_text, modelo="gpt-3.5-turbo"):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": modelo,
        "messages": [{"role": "user", "content": prompt_text}],
        "temperature": 0.7
    }
    try:
        for tentativa in range(3):
            try:
                response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data, timeout=60)
                response.raise_for_status() 
                return response.json()['choices'][0]['message']['content']
            except requests.exceptions.HTTPError as e:
                if response.status_code == 429:
                    time.sleep(2 * (tentativa + 1))
                    continue
                return f"Erro na API (HTTP {response.status_code}): {response.text}"
            except Exception:
                if tentativa == 2: raise
                time.sleep(1)
    except Exception as e:
        return f"Erro fatal na conexão: {e}"

# ==============================================================================
# FUNÇÕES AUXILIARES (UNIFICADOR DE DADOS)
# ==============================================================================
def processar_cpfs(df):
    df = df.dropna(how='all', axis=1)
    df.columns = [str(c).strip() for c in df.columns] 
    
    if 'CPF' not in df.columns:
        st.error("Erro: Coluna 'CPF' não encontrada na origem.")
        return None
    
    nome_coluna_corte = "Pergunta" 
    if nome_coluna_corte not in df.columns:
        st.error(f"Erro: Coluna '{nome_coluna_corte}' não encontrada.")
        return None

    try:
        idx_corte = df.columns.get_loc(nome_coluna_corte)
    except KeyError:
        st.error("Erro ao localizar índice da coluna de corte.")
        return None

    dados_processados = []
    grouped = df.groupby('CPF', sort=False)

    for cpf, group in grouped:
        if pd.isna(cpf) or str(cpf).strip() == "":
            continue

        linha_base = group.iloc[0].to_dict()
        
        if len(group) > 1:
            segunda_linha = group.iloc[1]
            colunas_extras = df.columns[idx_corte:]
            
            for col in colunas_extras:
                novo_nome = f"{col}_2"
                linha_base[novo_nome] = segunda_linha[col]
        
        dados_processados.append(linha_base)

    return pd.DataFrame(dados_processados)


# ==============================================================================
# INTERFACE PRINCIPAL
# ==============================================================================

st.title("🚀 Sistema Integrado de Relatórios")

tab1, tab2 = st.tabs(["📂 1. Unificador de Dados (ETL)", "🤖 2. Gerador com GPT"])

# ------------------------------------------------------------------------------
# ABA 1: UNIFICADOR DE DADOS
# ------------------------------------------------------------------------------
with tab1:
    st.header("Limpeza e Unificação para Google Sheets")
    st.markdown(f"""
    **Status da Configuração:**
    - **Origem (Fixa):** [Planilha de Dados Brutos]({URL_DADOS_BRUTOS})
    - **Destino (Fixa):** [Planilha Mestra]({URL_G_SHEET_LINK})
    """)
    
    if st.button("🔄 Carregar Dados da Origem", type="secondary"):
        st.session_state.dados_carregados = False 
        try:
            conn = st.connection("gsheets", type=GSheetsConnection)

            with st.spinner("Lendo dados brutos..."):
                df_bruto = conn.read(spreadsheet=URL_DADOS_BRUTOS, header=6, ttl=0)
                
                if len(df_bruto) > 7:
                    df_limpo = df_bruto.iloc[7:].reset_index(drop=True)
                    st.toast("Linhas de teste removidas.", icon="ℹ️")
                else:
                    df_limpo = df_bruto

                st.session_state.df_limpo_cache = df_limpo
                st.session_state.dados_carregados = True
                
        except Exception as e:
            st.error(f"Erro na leitura da origem: {e}")

    if st.session_state.get('dados_carregados') and 'df_limpo_cache' in st.session_state:
        df_display = st.session_state.df_limpo_cache
        
        with st.expander("Ver Prévia dos Dados Brutos Carregados", expanded=True):
            st.dataframe(df_display, use_container_width=True, height=200)

        st.divider()

        if st.button("⚙️ Processar e Atualizar Planilha Mestra", type="primary"):
            if df_display.empty:
                st.error("Não há dados para processar.")
            else:
                df_filtrado_unif = processar_cpfs(df_display)
                
                if df_filtrado_unif is not None and not df_filtrado_unif.empty:
                    st.info(f"Processado! {len(df_filtrado_unif)} registros únicos encontrados. Gravando...")
                    
                    try:
                        conn = st.connection("gsheets", type=GSheetsConnection)
                        conn.update(spreadsheet=URL_G_SHEET_LINK, data=df_filtrado_unif)
                        
                        st.success("✅ Sucesso! Dados atualizados na Planilha Mestra.")
                        st.markdown(f"**[Clique aqui para conferir a Planilha Mestra]({URL_G_SHEET_LINK})**")
                        st.dataframe(df_filtrado_unif, use_container_width=True)
                        
                    except Exception as e_write:
                        st.error(f"Erro ao escrever na planilha: {e_write}")
                else:
                    st.warning("Erro no processamento dos dados.")

# ------------------------------------------------------------------------------
# ABA 2: GERADOR GPT
# ------------------------------------------------------------------------------
with tab2:
    st.header("Automação de Relatórios com IA")
    
    if 'processamento_concluido' not in st.session_state:
        st.session_state.processamento_concluido = False
    if 'zip_prompts' not in st.session_state:
        st.session_state.zip_prompts = None
    if 'todos_prompts' not in st.session_state:
        st.session_state.todos_prompts = [] 
    if 'respostas_geradas' not in st.session_state:
        st.session_state.respostas_geradas = [] 
    if 'zip_respostas' not in st.session_state:
        st.session_state.zip_respostas = None

    col_config, col_main = st.columns([1, 3])

    with col_config:
        st.subheader("Configurações IA")
        api_key = st.text_input("OpenAI API Key", type="password", key="gpt_api_key")
        modelo_gpt = st.selectbox("Modelo GPT", ["gpt-3.5-turbo", "gpt-4", "gpt-4o", "gpt-5.2"], key="gpt_model_select")
        
        st.divider()
        st.subheader("Arquivos")
        uploaded_excel = st.file_uploader("Upload Excel (Opcional)", type=["xlsx"], key="gpt_excel_upload")
        uploaded_template = st.file_uploader("Upload Template (Opcional)", type=["txt"], key="gpt_template_upload")

    with col_main:
        df_gpt = None
        template_content = ""

        try:
            if uploaded_excel:
                df_gpt = pd.read_excel(uploaded_excel, sheet_name="dupar")
            else:
                try:
                    df_gpt = pd.read_excel(URL_EXPORT_EXCEL, sheet_name="dupar")
                    st.success(f"✅ Usando dados da Planilha Mestra")
                except Exception as e:
                    st.warning(f"Não foi possível ler a Planilha Mestra automaticamente. Use a Aba 1 primeiro. ({e})")

            if uploaded_template:
                template_content = uploaded_template.read().decode("utf-8")
                st.info("Usando template carregado pelo usuário.")
            else:
                template_content = TEMPLATE_PADRAO
                st.info("Usando template padrão embutido.")

        except Exception as e:
            st.error(f"Erro ao carregar arquivos: {e}")

        if df_gpt is not None and template_content:
            c1, c2 = st.columns([1, 1])
            
            with c1:
                st.markdown("##### Seleção de Colunas")
                todas_colunas = df_gpt.columns.tolist()
                
                colunas_selecionadas = st.multiselect(
                    "Colunas para o prompt:",
                    options=todas_colunas,
                    default=todas_colunas,
                    key="cols_multiselect"
                )
                st.caption(f"Registros: {len(df_gpt)}")

                if st.button("📝 Preparar Prompts", type="primary", key="btn_prep_prompts"):
                    if not colunas_selecionadas:
                        st.error("Selecione colunas.")
                    else:
                        zip_buffer = io.BytesIO()
                        prompts_gerados = [] 

                        with zipfile.ZipFile(zip_buffer, "w") as zf:
                            bar = st.progress(0)
                            for i, row in df_gpt.iterrows():
                                dados_filtrados = {col: row[col] for col in colunas_selecionadas}
                                conteudo_prompt = ler_e_substituir_template(template_content, dados_filtrados)
                                
                                nome_pessoa = str(row.get('Nome', f'Registro {i+1}'))
                                nome_limpo = re.sub(r'[^a-zA-Z0-9\s]', '', nome_pessoa).replace(' ', '_')
                                
                                zf.writestr(f"{nome_limpo}_prompt.txt", conteudo_prompt)
                                prompts_gerados.append({
                                    "id": i,
                                    "nome": nome_pessoa,
                                    "nome_arquivo": nome_limpo,
                                    "conteudo": conteudo_prompt
                                })
                                bar.progress((i + 1) / len(df_gpt))
                        
                        zip_buffer.seek(0)
                        st.session_state.processamento_concluido = True
                        st.session_state.zip_prompts = zip_buffer
                        st.session_state.todos_prompts = prompts_gerados
                        st.session_state.respostas_geradas = [] 
                        st.session_state.zip_respostas = None
                        st.rerun()

            with c2:
                st.markdown("##### Visualização")
                if colunas_selecionadas:
                    st.dataframe(df_gpt[colunas_selecionadas], height=250, hide_index=True)
                else:
                    st.dataframe(df_gpt, height=250, hide_index=True)

            if st.session_state.processamento_concluido:
                st.divider()
                cp, cg = st.columns([1, 1])
                
                with cp:
                    st.markdown("##### 1. Prompts Prontos")
                    st.download_button("⬇️ Baixar Prompts (.zip)", st.session_state.zip_prompts, "prompts.zip", "application/zip")
                    
                    with st.container(height=500):
                        for item in st.session_state.todos_prompts:
                            with st.expander(f"📄 {item['nome']}"):
                                st.text_area("Conteúdo", item['conteudo'], height=150, key=f"t_{item['id']}")

                with cg:
                    st.markdown("##### 2. Respostas IA")
                    
                    # --- NOVO: MULTISELECT PARA FILTRAR ---
                    opcoes_nomes = [p['nome'] for p in st.session_state.todos_prompts]
                    
                    selecionados_gpt = st.multiselect(
                        "Selecione quem você deseja enviar:",
                        options=opcoes_nomes,
                        default=opcoes_nomes, # Padrão: Seleciona todos
                        help="Remova nomes da lista se não quiser enviá-los."
                    )
                    
                    if st.button(f"🚀 Processar Selecionados ({len(selecionados_gpt)})", type="primary", key="btn_run_gpt"):
                        if not api_key:
                            st.warning("⚠️ Insira a API Key.")
                        elif not selecionados_gpt:
                            st.warning("⚠️ Selecione pelo menos um registro na caixa acima.")
                        else:
                            st.session_state.respostas_geradas = [] 
                            status_text = st.empty()
                            bar_gpt = st.progress(0)
                            
                            # Filtra a lista completa com base na seleção
                            fila_processamento = [p for p in st.session_state.todos_prompts if p['nome'] in selecionados_gpt]
                            
                            zip_resp_io = io.BytesIO()
                            with zipfile.ZipFile(zip_resp_io, "w") as zf_resp:
                                total = len(fila_processamento)
                                for idx, item in enumerate(fila_processamento):
                                    status_text.text(f"Processando {idx+1}/{total}: {item['nome']}...")
                                    
                                    resp_txt = chamar_gpt(api_key, item['conteudo'], modelo_gpt)
                                    
                                    st.session_state.respostas_geradas.append({"nome": item['nome'], "resposta": resp_txt})
                                    
                                    docx_bytes = criar_docx_bytes(resp_txt)
                                    zf_resp.writestr(f"RESPOSTA_{item['nome_arquivo']}.docx", docx_bytes)
                                    
                                    bar_gpt.progress((idx + 1) / total)
                            
                            zip_resp_io.seek(0)
                            st.session_state.zip_respostas = zip_resp_io
                            status_text.success("Concluído!")
                            st.rerun()

                    if st.session_state.zip_respostas:
                        st.download_button("⬇️ Baixar Respostas (.zip)", st.session_state.zip_respostas, "respostas_docx.zip", "application/zip", type="primary")
                    
                    with st.container(height=500):
                        if not st.session_state.respostas_geradas:
                            st.info("Aguardando processamento...")
                        else:
                            for r in st.session_state.respostas_geradas:
                                icon = "❌" if "Erro" in r['resposta'] else "✅"
                                with st.expander(f"{icon} {r['nome']}"):
                                    st.markdown(r['resposta'])

        elif df_gpt is None:
            st.info("Aguardando dados (Verifique Aba 1)...")